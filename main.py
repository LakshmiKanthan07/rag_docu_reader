import os
import json
import logging
import tempfile
import asyncio
from collections import OrderedDict
from contextlib import asynccontextmanager
from datetime import datetime, timedelta, timezone
from typing import Dict, List, Optional
import threading

from dotenv import load_dotenv
load_dotenv()

# ── Structured logging ────────────────────────────────────────────────────────
class JSONFormatter(logging.Formatter):
    def format(self, record: logging.LogRecord) -> str:
        return json.dumps({
            "ts":      self.formatTime(record, "%Y-%m-%dT%H:%M:%S"),
            "level":   record.levelname,
            "logger":  record.name,
            "msg":     record.getMessage(),
            **({"exc": self.formatException(record.exc_info)} if record.exc_info else {}),
        })

handler = logging.StreamHandler()
handler.setFormatter(JSONFormatter())
logging.basicConfig(level=logging.INFO, handlers=[handler])
logger = logging.getLogger(__name__)

from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, field_validator
import uuid

from langchain_ollama import OllamaEmbeddings
from langchain_groq import ChatGroq
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader, JSONLoader
from langchain_core.documents.base import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
import docx
import openpyxl

from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ── Config ────────────────────────────────────────────────────────────────────
MAX_FILE_SIZE     = int(os.getenv("MAX_FILE_SIZE_MB",  "50"))   * 1024 * 1024
MAX_FILES         = int(os.getenv("MAX_FILES_PER_UPLOAD", "0"))   # 0 = no hard limit
CHUNK_SIZE        = int(os.getenv("CHUNK_SIZE",        "500"))
CHUNK_OVERLAP     = int(os.getenv("CHUNK_OVERLAP",     "50"))
TOP_K             = int(os.getenv("TOP_K",             "5"))
SESSION_TTL_HOURS = int(os.getenv("SESSION_TTL_HOURS", "24"))
PERSIST_DIR       = os.getenv("FAISS_PERSIST_DIR",     "./faiss_store")
API_KEY           = os.getenv("API_KEY",               "")
VS_CACHE_SIZE     = int(os.getenv("VS_CACHE_SIZE",     "20"))

ALLOWED_EXTENSIONS = {
    ".pdf", ".csv", ".json",
    ".txt", ".md", ".html", ".htm", ".log", ".xml",
    ".yaml", ".yml", ".toml", ".ini", ".cfg",
    ".py", ".java", ".js", ".ts", ".css", ".scss",
    ".docx", ".xlsx",
}
TEXT_EXTENSIONS = {
    ".txt", ".md", ".html", ".htm", ".log", ".xml",
    ".yaml", ".yml", ".toml", ".ini", ".cfg",
    ".py", ".java", ".js", ".ts", ".css", ".scss",
}

os.makedirs(PERSIST_DIR, exist_ok=True)

# ── Global metrics counters ───────────────────────────────────────────────────
_metrics: Dict[str, int] = {
    "uploads_total":   0,
    "uploads_failed":  0,
    "questions_total": 0,
    "sessions_active": 0,
}

# ── Rate limiter ──────────────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address)

# ── Lifespan (replaces deprecated @app.on_event) ─────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    # startup
    if not os.getenv("GROQ_API_KEY"):
        logger.error("GROQ_API_KEY is not set — LLM calls will fail.")
    try:
        embeddings.embed_query("ping")
        logger.info("Ollama embeddings reachable.")
    except Exception as exc:
        logger.error("Ollama unreachable at startup: %s", exc)
    yield
    # shutdown
    sessions.shutdown()
    logger.info("SessionStore shut down cleanly.")

# ── App ───────────────────────────────────────────────────────────────────────
app = FastAPI(
    title="DocTalk API",
    description="Upload files and chat with their content — streaming, memory, persistent storage.",
    version="2.1.0",
    lifespan=lifespan,
)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

origins = os.getenv(
    "CORS_ORIGINS",
    "http://localhost:8000,http://127.0.0.1:8000,http://localhost:5500,http://127.0.0.1:5500",
).split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "DELETE"],
    allow_headers=["*"],
)

# ── Auth dependency ───────────────────────────────────────────────────────────
def require_api_key(request: Request) -> None:
    if not API_KEY:
        return
    if request.headers.get("x-api-key", "") != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing API key.")

# ── LLM & Embeddings ──────────────────────────────────────────────────────────
embeddings = OllamaEmbeddings(model=os.getenv("EMBED_MODEL", "nomic-embed-text"))

llm = ChatGroq(
    model=os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
    api_key=os.getenv("GROQ_API_KEY"),
    temperature=float(os.getenv("TEMPERATURE", "0.3")),
    max_tokens=int(os.getenv("MAX_TOKENS", "1024")),
)

# ── RAG prompt ────────────────────────────────────────────────────────────────
RAG_PROMPT = ChatPromptTemplate.from_template(
    """You are a precise document assistant. Answer questions strictly from the context below.

Rules:
- Use ONLY the provided context. No outside knowledge.
- If the answer is absent from the context, say: "I don't know based on the provided document."
- Be concise and direct. Skip filler like "Based on the context...".
- After your answer, list the sources you used on a new line in the format:
  Sources: [page X], [page Y]  (use the page numbers shown in the context blocks)
  If no page numbers are available, omit the Sources line.
- Never guess, infer beyond the text, or hallucinate.

--- DOCUMENT CONTEXT ---
{context}
--- END CONTEXT ---

Conversation history:
{history}

User: {question}
Assistant:"""
)

# ── LRU vectorstore cache ─────────────────────────────────────────────────────
class LRUVectorstoreCache:
    """
    Thread-safe in-memory LRU cache for FAISS vectorstores.
    Prevents repeated disk I/O on every query for the same session.
    Max size is VS_CACHE_SIZE; evicted entries are dropped from RAM
    but remain on disk for reload if needed.
    """

    def __init__(self, maxsize: int = 20) -> None:
        self._cache: OrderedDict[str, FAISS] = OrderedDict()
        self._maxsize = maxsize
        self._lock = threading.Lock()

    def get(self, key: str) -> Optional[FAISS]:
        with self._lock:
            if key not in self._cache:
                return None
            self._cache.move_to_end(key)
            return self._cache[key]

    def put(self, key: str, vs: FAISS) -> None:
        with self._lock:
            if key in self._cache:
                self._cache.move_to_end(key)
            else:
                if len(self._cache) >= self._maxsize:
                    self._cache.popitem(last=False)  # evict LRU
            self._cache[key] = vs

    def remove(self, key: str) -> None:
        with self._lock:
            self._cache.pop(key, None)


vs_cache = LRUVectorstoreCache(maxsize=VS_CACHE_SIZE)

# ── Session store ─────────────────────────────────────────────────────────────
class SessionStore:
    """
    Registry of active sessions.
    Vectorstores are persisted to disk; hot sessions are cached in RAM via LRUVectorstoreCache.
    Chat history is kept in memory (last 20 turns per session).
    """

    def __init__(self, ttl_hours: int = 24) -> None:
        self._meta: Dict[str, dict] = {}
        self._ttl  = timedelta(hours=ttl_hours)
        self._lock = threading.Lock()
        self._stop = threading.Event()
        self._start_cleanup_thread()
        logger.info("SessionStore ready — ttl_h=%d dir=%s", ttl_hours, PERSIST_DIR)

    # ── public ────────────────────────────────────────────────

    def create(self, session_id: str, vectorstore: FAISS) -> None:
        vectorstore.save_local(self._path(session_id))
        vs_cache.put(session_id, vectorstore)
        with self._lock:
            self._meta[session_id] = {"created_at": datetime.now(), "history": []}
        _metrics["sessions_active"] = len(self._meta)
        logger.info("Session created: %s", session_id[:8])

    def get_vectorstore(self, session_id: str) -> FAISS:
        self._validate(session_id)
        vs = vs_cache.get(session_id)
        if vs is not None:
            return vs
        path = self._path(session_id)
        if not os.path.isdir(path):
            raise HTTPException(status_code=404, detail="Session data missing on disk.")
        vs = FAISS.load_local(path, embeddings, allow_dangerous_deserialization=True)
        vs_cache.put(session_id, vs)
        return vs

    def get_history(self, session_id: str) -> List[dict]:
        self._validate(session_id)
        with self._lock:
            return list(self._meta[session_id]["history"])

    def append_history(self, session_id: str, role: str, content: str) -> None:
        """
        Append a turn to history only if the session still exists.
        Silently no-ops if the session was removed (e.g. expired mid-stream).
        """
        with self._lock:
            meta = self._meta.get(session_id)
            if meta is None:
                return
            hist = meta["history"]
            hist.append({"role": role, "content": content})
            meta["history"] = hist[-20:]

    def remove(self, session_id: str) -> None:
        vs_cache.remove(session_id)
        with self._lock:
            self._meta.pop(session_id, None)
        import shutil as _shutil
        path = self._path(session_id)
        if os.path.isdir(path):
            _shutil.rmtree(path, ignore_errors=True)
        _metrics["sessions_active"] = len(self._meta)
        logger.info("Session removed: %s", session_id[:8])

    def active_count(self) -> int:
        with self._lock:
            return len(self._meta)

    def shutdown(self) -> None:
        self._stop.set()

    # ── private ───────────────────────────────────────────────

    def _path(self, session_id: str) -> str:
        return os.path.join(PERSIST_DIR, session_id)

    def _validate(self, session_id: str) -> None:
        with self._lock:
            meta = self._meta.get(session_id)
        if meta is None:
            raise HTTPException(
                status_code=404,
                detail="Session not found or expired. Please upload the file again.",
            )
        if datetime.now() - meta["created_at"] >= self._ttl:
            self.remove(session_id)
            raise HTTPException(
                status_code=404,
                detail="Session expired. Please upload the file again.",
            )

    def _cleanup(self) -> None:
        with self._lock:
            now = datetime.now()
            expired = [sid for sid, m in self._meta.items() if now - m["created_at"] >= self._ttl]
        for sid in expired:
            self.remove(sid)
        if expired:
            logger.info("Cleaned up %d expired session(s).", len(expired))

    def _start_cleanup_thread(self) -> None:
        def loop():
            while not self._stop.wait(3600):
                self._cleanup()
        threading.Thread(target=loop, daemon=True).start()


sessions = SessionStore(ttl_hours=SESSION_TTL_HOURS)

# ── Document loaders ──────────────────────────────────────────────────────────
def load_text_file(path: str) -> List[Document]:
    return TextLoader(path, encoding="utf-8").load()

def load_docx_file(path: str) -> List[Document]:
    d = docx.Document(path)
    parts: List[str] = []
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    content = "\n\n".join(parts).strip()
    return [Document(page_content=content)] if content else []

def load_xlsx_file(path: str) -> List[Document]:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts: List[str] = []
    for sheet in wb.worksheets:
        rows = ["\t".join("" if v is None else str(v) for v in row)
                for row in sheet.iter_rows(values_only=True)]
        if any(r.strip() for r in rows):
            parts.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
    content = "\n\n".join(parts).strip()
    return [Document(page_content=content)] if content else []

LOADER_MAP = {
    ".pdf":  lambda p: PyPDFLoader(p).load(),
    ".csv":  lambda p: CSVLoader(p).load(),
    ".json": lambda p: JSONLoader(p, jq_schema=".", text_content=False).load(),
    ".docx": load_docx_file,
    ".xlsx": load_xlsx_file,
    **{ext: load_text_file for ext in TEXT_EXTENSIONS},
}

def load_documents(file_path: str, filename: str) -> List[Document]:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{ext}'. Allowed: {sorted(ALLOWED_EXTENSIONS)}",
        )
    loader = LOADER_MAP.get(ext)
    if loader is None:
        raise HTTPException(status_code=415, detail=f"No loader configured for '{ext}'.")
    return loader(file_path)

def cleanup_file(path: str) -> None:
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception as exc:
        logger.warning("Could not delete temp file %s: %s", path, exc)

# ── /upload  (multi-file → single merged session) ────────────────────────────
@app.post("/upload", dependencies=[Depends(require_api_key)])
@limiter.limit("10/minute")
async def upload_files(request: Request, files: List[UploadFile] = File(...)) -> dict:
    if not files:
        raise HTTPException(status_code=400, detail="No files provided.")

    if MAX_FILES and len(files) > MAX_FILES:
        raise HTTPException(
            status_code=400,
            detail=f"Too many files. Maximum allowed per upload is {MAX_FILES}.",
        )

    session_id  = str(uuid.uuid4())
    all_chunks: List[Document] = []
    tmp_paths:  List[str]      = []
    file_results               = []   # per-file outcome summary

    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    for file in files:
        if not file.filename:
            file_results.append({"file": "(unnamed)", "status": "skipped", "reason": "no filename"})
            continue

        ext = os.path.splitext(file.filename)[1].lower() or ".txt"
        if ext not in ALLOWED_EXTENSIONS:
            _metrics["uploads_failed"] += 1
            file_results.append({
                "file": file.filename, "status": "skipped",
                "reason": f"unsupported type '{ext}'",
            })
            continue

        contents = await file.read()
        if len(contents) > MAX_FILE_SIZE:
            _metrics["uploads_failed"] += 1
            file_results.append({
                "file": file.filename, "status": "skipped",
                "reason": f"exceeds {MAX_FILE_SIZE // (1024*1024)} MB limit",
            })
            continue

        tmp_path: Optional[str] = None
        try:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(contents)
                tmp_path = tmp.name
            tmp_paths.append(tmp_path)

            docs = load_documents(tmp_path, file.filename)
            if not docs:
                file_results.append({"file": file.filename, "status": "skipped", "reason": "no text extracted"})
                continue

            # Tag every chunk with its source filename so citations stay traceable
            for i, doc in enumerate(docs):
                doc.metadata.setdefault("source", file.filename)
                if "page" not in doc.metadata:
                    doc.metadata["page"] = i + 1

            chunks = splitter.split_documents(docs)
            if not chunks:
                file_results.append({"file": file.filename, "status": "skipped", "reason": "empty after splitting"})
                continue

            all_chunks.extend(chunks)
            file_results.append({"file": file.filename, "status": "ok", "chunks": len(chunks)})
            logger.info("Loaded — file=%s chunks=%d", file.filename, len(chunks))

        except HTTPException as exc:
            _metrics["uploads_failed"] += 1
            file_results.append({"file": file.filename, "status": "error", "reason": exc.detail})
        except Exception as exc:
            _metrics["uploads_failed"] += 1
            logger.exception("Load error for file: %s", file.filename)
            file_results.append({"file": file.filename, "status": "error", "reason": str(exc)})

    # Clean up all temp files regardless of outcome
    for p in tmp_paths:
        cleanup_file(p)

    if not all_chunks:
        raise HTTPException(
            status_code=422,
            detail={"message": "No usable content extracted from any file.", "files": file_results},
        )

    try:
        # Embed all chunks from all files into one shared vectorstore
        vectorstore = await asyncio.to_thread(FAISS.from_documents, all_chunks, embeddings)
        sessions.create(session_id, vectorstore)
    except Exception as exc:
        _metrics["uploads_failed"] += 1
        logger.exception("Vectorstore creation failed — session=%s", session_id[:8])
        raise HTTPException(status_code=500, detail=f"Indexing error: {exc}")

    ok_count = sum(1 for r in file_results if r["status"] == "ok")
    _metrics["uploads_total"] += ok_count
    logger.info(
        "Upload OK — session=%s files_ok=%d total_chunks=%d",
        session_id[:8], ok_count, len(all_chunks),
    )
    return {
        "message":      f"{ok_count} file(s) processed and merged into one session.",
        "session_id":   session_id,
        "total_chunks": len(all_chunks),
        "files":        file_results,
    }

# ── /ask  (streaming) ─────────────────────────────────────────────────────────
class Query(BaseModel):
    question: str
    session_id: str

    @field_validator("question")
    @classmethod
    def not_empty(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("Question cannot be empty.")
        if len(v) > 2000:
            raise ValueError("Question too long (max 2000 characters).")
        return v


def _fmt_history(history: List[dict]) -> str:
    if not history:
        return "(no prior conversation)"
    lines = []
    for h in history:
        speaker = "User" if h["role"] == "human" else "Assistant"
        lines.append(f"{speaker}: {h['content']}")
    return "\n".join(lines)


def _fmt_context(docs: List[Document]) -> str:
    parts = []
    for doc in docs:
        page = doc.metadata.get("page", "?")
        parts.append(f"[page {page}]\n{doc.page_content}")
    return "\n\n".join(parts)


@app.post("/ask", dependencies=[Depends(require_api_key)])
@limiter.limit("30/minute")
async def ask_question(request: Request, query: Query) -> StreamingResponse:
    logger.info("Question — session=%s q=%r", query.session_id[:8], query.question)
    _metrics["questions_total"] += 1

    vectorstore = sessions.get_vectorstore(query.session_id)
    history     = sessions.get_history(query.session_id)

    retriever = vectorstore.as_retriever(search_kwargs={"k": TOP_K})
    docs      = await asyncio.to_thread(retriever.invoke, query.question)
    context   = _fmt_context(docs)

    final_prompt = RAG_PROMPT.invoke({
        "context":  context,
        "history":  _fmt_history(history),
        "question": query.question,
    })

    # FIX: buffer both turns and write them together only after the stream
    # completes. This prevents a partial/empty assistant turn being recorded
    # if the client disconnects mid-stream, and keeps history consistent.
    async def stream_tokens():
        collected: List[str] = []
        try:
            async for chunk in llm.astream(final_prompt):
                token = chunk.content
                if token:
                    collected.append(token)
                    yield token
        except Exception as exc:
            logger.exception("Streaming error — session=%s", query.session_id[:8])
            yield f"\n[ERROR: {exc}]"
        finally:
            full_answer = "".join(collected)
            # Append both turns atomically only after streaming finishes.
            # append_history silently no-ops if the session expired mid-stream.
            sessions.append_history(query.session_id, "human",     query.question)
            sessions.append_history(query.session_id, "assistant", full_answer)

    return StreamingResponse(stream_tokens(), media_type="text/plain")


# ── /session/{id}  DELETE ─────────────────────────────────────────────────────
@app.delete("/session/{session_id}", dependencies=[Depends(require_api_key)])
async def delete_session(session_id: str) -> dict:
    sessions.remove(session_id)
    return {"message": "Session deleted."}


# ── /health ───────────────────────────────────────────────────────────────────
@app.get("/health")
async def health() -> dict:
    # FIX: datetime.timezone is not callable; use datetime.now(timezone.utc)
    return {
        "status":    "ok",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "sessions":  sessions.active_count(),
    }


# ── /metrics ──────────────────────────────────────────────────────────────────
@app.get("/metrics")
async def metrics() -> dict:
    # FIX: same datetime.timezone bug corrected here too
    return {
        **_metrics,
        "sessions_active": sessions.active_count(),
        "vs_cache_size":   VS_CACHE_SIZE,
        "timestamp":       datetime.now(timezone.utc).isoformat(),
    }